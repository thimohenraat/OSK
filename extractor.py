from docx import Document
import fitz
import logging
import re

def extract_text_from_docx(docx_file_path):
    try:
        document = Document(docx_file_path)
        paragraph_texts = [paragraph.text for paragraph in document.paragraphs]
                # Voeg ook tekst uit tabellen toe
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    paragraph_texts.append(cell.text)

        full_text = "\n".join(paragraph_texts)
        return full_text
    except Exception as error:
        logging.error(f"Failed to extract text from DOCX at {docx_file_path}: {error}")
        return ""

def extract_text_from_pdf(pdf_path):
    try:
        with fitz.open(pdf_path) as pdf_document:
            page_texts = []
            for page in pdf_document:
                page_text = page.get_text("text")  # Haalt platte tekst op
                page_texts.append((page.number + 1, page_text))
            return page_texts
    except Exception as error:
        logging.error(f"Failed to extract text from PDF at {pdf_path}: {error}")
        return []

